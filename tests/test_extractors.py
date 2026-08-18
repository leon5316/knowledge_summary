import json

import pytest

from knowledge_summary.config import load_config
from knowledge_summary.extractors import extract_file
from knowledge_summary.extractors.base import ExtractorError

CFG = load_config()


def _extract(tmp_path, name, content, mode="w"):
    f = tmp_path / name
    if mode == "wb":
        f.write_bytes(content)
    else:
        f.write_text(content, encoding="utf-8")
    return extract_file(f, name, CFG)


def test_markdown_headings(tmp_path):
    doc = _extract(tmp_path, "doc.md", "# Title\n\ntext here\n\n## Sub\n\nmore\n")
    assert doc.format == "md"
    assert "text here" in doc.text
    kinds = [s.kind for s in doc.sections]
    assert kinds.count("heading") == 2


def test_python_sections(tmp_path):
    src = '"""mod doc"""\n\nclass Foo:\n    def bar(self):\n        pass\n\ndef baz():\n    pass\n'
    doc = _extract(tmp_path, "mod.py", src)
    assert doc.format == "py"
    titles = {s.title for s in doc.sections}
    assert {"Foo", "bar", "baz"} <= titles


def test_sql_sections(tmp_path):
    sql = "CREATE TABLE t (id INT);\n\nSELECT * FROM t;\n"
    doc = _extract(tmp_path, "schema.sql", sql)
    assert doc.meta["statements"] == 2


def test_csv(tmp_path):
    doc = _extract(tmp_path, "data.csv", "name,age\nAlice,30\nBob,25\n")
    assert "Alice" in doc.text
    assert "name" in doc.text


def test_json(tmp_path):
    doc = _extract(tmp_path, "data.json", json.dumps({"a": 1, "b": [1, 2]}))
    assert "a" in doc.text


def test_ipynb(tmp_path):
    nb = {"cells": [
        {"cell_type": "markdown", "source": ["# Notebook title"]},
        {"cell_type": "code", "source": ["x = 1\n"]},
    ]}
    doc = _extract(tmp_path, "nb.ipynb", json.dumps(nb))
    assert "Notebook title" in doc.text
    assert "x = 1" in doc.text


def test_html(tmp_path):
    html = "<html><head><title>T</title></head><body><h1>H1</h1><p>para</p></body></html>"
    doc = _extract(tmp_path, "page.html", html)
    assert "H1" in doc.text
    assert "para" in doc.text


def test_docx(tmp_path):
    try:
        import docx
    except ImportError:
        pytest.skip("python-docx 未安装")
    p = tmp_path / "w.docx"
    d = docx.Document()
    d.add_heading("Doc Title", level=1)
    d.add_paragraph("Some body text")
    d.save(str(p))
    doc = extract_file(p, "w.docx", CFG)
    assert "Doc Title" in doc.text
    assert "Some body text" in doc.text
    assert any(s.kind == "heading" for s in doc.sections)


def test_unknown_extension_skipped_by_default(tmp_path):
    f = tmp_path / "data.xyz"
    f.write_text("hello", encoding="utf-8")
    with pytest.raises(ExtractorError):
        extract_file(f, "data.xyz", CFG)


def test_unknown_extension_as_text(tmp_path):
    cfg = load_config(cli_overrides={"general": {"treat_unknown_as_text": True}})
    f = tmp_path / "data.xyz"
    f.write_text("hello", encoding="utf-8")
    doc = extract_file(f, "data.xyz", cfg)
    assert "hello" in doc.text


def test_pdf_empty_content_error(tmp_path):
    try:
        from pypdf import PdfWriter
    except ImportError:
        pytest.skip("pypdf 未安装")
    p = tmp_path / "blank.pdf"
    w = PdfWriter()
    w.add_blank_page(width=200, height=200)
    with open(p, "wb") as fh:
        w.write(fh)
    with pytest.raises(ExtractorError):
        extract_file(p, "blank.pdf", CFG)


def test_doc_missing_tool_error(tmp_path):
    f = tmp_path / "old.doc"
    f.write_bytes(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1fake")
    with pytest.raises(ExtractorError, match="antiword|LibreOffice"):
        extract_file(f, "old.doc", CFG)
