"""Word .docx 解析器：段落（含标题样式）+ 表格。"""
from __future__ import annotations

from pathlib import Path

from ..models import Document, Section, Span
from .base import Extractor, ExtractorError


class DocxExtractor(Extractor):
    extensions = {"docx"}
    display_name = "docx"

    def extract(self, path: Path, rel_path: str, cfg) -> Document:
        try:
            import docx  # python-docx
        except ImportError as e:
            raise ExtractorError("解析 .docx 需要 python-docx，请先 pip install python-docx") from e

        document = docx.Document(str(path))
        lines = []
        sections = []
        line_no = 1

        for para in document.paragraphs:
            style = (para.style.name or "") if para.style else ""
            text = para.text.strip()
            if not text:
                line_no += 1
                continue
            lines.append(text)
            if style and ("Heading" in style or style.startswith("标题")):
                sections.append(Section(
                    title=text[:80], kind="heading",
                    span=Span(str(path), rel_path, line_no, line_no),
                    meta={"style": style},
                ))
            line_no += 1

        for table in document.tables:
            rows = []
            for row in table.rows:
                cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                rows.append(" | ".join(cells))
            block = "\n".join(rows)
            if block.strip():
                lines.append(block)
                sections.append(Section(
                    title=f"Table ({len(table.rows)} rows)", kind="table",
                    span=Span(str(path), rel_path, line_no, line_no + len(rows)),
                ))
                line_no += len(rows)

        text = "\n\n".join(lines)
        return self.make_doc(path, rel_path, "docx", text,
                             sections=sections[:2000],
                             meta={"paragraphs": len(document.paragraphs),
                                   "tables": len(document.tables),
                                   "note": "Line numbers are approximate sequential indexes"})
